#!/usr/bin/env python3
"""
recon_report_docx.py — build a .docx exam report from a recon_scan.py
project, matching the structure of OffSec's official OSCP exam report
template (OffSec Certified Professional Exam Report, v2.0):

    1. OffSec Certified Professional Exam Report
       1.1 Introduction / 1.2 Objective / 1.3 Requirements
    2. High-Level Summary / 2.1 Recommendations
    3. Methodologies (Information Gathering / Service Enumeration /
       Penetration / Maintaining Access / House Cleaning)
    4. Independent Challenges (standalone, non-domain-joined hosts)
    5. Active Directory Set (domain-joined hosts)

Usage:
    python3 recon_report_docx.py <project_dir> [-o OUTPUT.docx] [--all]
        [--candidate NAME] [--osid OSID] [--email EMAIL]

Hosts are bucketed into "Independent Challenges" vs "Active Directory
Set" based on the Domain / Workgroup field in each host's System
Information note section — domain-joined hosts (anything other than
blank/WORKGROUP) go under Active Directory Set.

What's auto-filled: host list, port/service enumeration, credentials
(from the manifest), the technical attack narrative and flag hashes
(from notes), tunneling detail. What's deliberately left as a
placeholder for you to write: Vulnerability Explanation/Fix/Severity and
the "steps to reproduce" summary for each finding, and the High-Level
Summary/Recommendations narrative — those are exactly the analysis and
understanding OSCP is grading, so this tool won't fabricate them for
you. Screenshots are always a placeholder too, obviously.
"""
import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[-] python-docx not found. Install with: pip install python-docx --break-system-packages")
    sys.exit(1)

try:
    from recon_common import load_manifest
    from recon_report import parse_note_sections, load_note, access_info
except ImportError as e:
    print(f"[-] Could not import recon_common.py / recon_report.py: {e}")
    print("    Keep recon_report_docx.py in the same directory as the rest of the project.")
    sys.exit(1)

H3_RE = re.compile(r'(?m)^### (.+)$')


# ── Note parsing helpers ────────────────────────────────────────────────────
def parse_h3_subsections(text: str) -> dict:
    parts = H3_RE.split(text)
    sections = {}
    it = iter(parts[1:])
    for heading, body in zip(it, it):
        sections[heading.strip()] = body.strip('\n')
    return sections


def parse_network_access(section_text: str) -> dict:
    result = {'access': 'Direct', 'pivot_tool': '', 'pivot_host': '', 'commands': ''}
    m = re.search(r'\|\s*\*\*Access\*\*\s*\|\s*(.*?)\s*\|', section_text)
    if m:
        result['access'] = m.group(1).strip() or 'Direct'
    m = re.search(r'\|\s*\*\*Pivot Tool\*\*\s*\|\s*(.*?)\s*\|', section_text)
    if m:
        result['pivot_tool'] = m.group(1).strip()
    m = re.search(r'\|\s*\*\*Pivot Via Host\*\*\s*\|\s*(.*?)\s*\|', section_text)
    if m:
        result['pivot_host'] = m.group(1).strip()
    m = re.search(r'```\s*\n(.*?)\n```', section_text, re.DOTALL)
    if m:
        result['commands'] = m.group(1).strip()
    return result


def is_pivoted(net_access: dict) -> bool:
    access = net_access.get('access', '').strip().lower()
    return bool(access) and access != 'direct'


def parse_flags(sections: dict) -> dict:
    flags_text = sections.get('Flags', '')
    sub = parse_h3_subsections(flags_text) if flags_text else {}
    result = {}
    for key, out_key in [('User Flag', 'local'), ('Root / Admin Flag', 'proof')]:
        body = sub.get(key, '')
        # findall (not a single search) so the first fence pair is picked
        # correctly even when that first block is empty — a plain
        # non-greedy search can otherwise pair the first opening fence
        # with a LATER block's closing fence (e.g. the "Proof command
        # output" block that follows), swallowing everything between.
        blocks = re.findall(r'```(.*?)```', body, re.DOTALL)
        hash_val = blocks[0].strip() if blocks else ''
        path_m = re.search(r'\*\*Path:\*\*\s*(.+)', body)
        result[out_key] = {'hash': hash_val, 'path': path_m.group(1).strip() if path_m else ''}
    return result


def parse_domain(sections: dict) -> str:
    """Pulls Domain / Workgroup from System Information — used to bucket
    the host into Independent Challenges vs Active Directory Set."""
    sys_info = sections.get('System Information', '')
    m = re.search(r'\|\s*\*\*Domain / Workgroup\*\*\s*\|\s*(.*?)\s*\|', sys_info)
    return m.group(1).strip() if m else ''


def is_ad_domain_joined(domain: str) -> bool:
    d = domain.strip().upper()
    return bool(d) and d != 'WORKGROUP'


# ── Generic markdown -> docx block renderer (see recon_report_docx history —
# fixes **bold**/`code`/```fences```/tables/checkboxes showing as literal text) ──
INLINE_TOKEN_RE = re.compile(r'(\*\*.+?\*\*|`[^`]+`)')
LABEL_ONLY_RE = re.compile(r'^\*\*(.+?):\*\*\s*$')
LABEL_PLACEHOLDER_RE = re.compile(r'^\*\*(.+?):\*\*\s*_\(.*\)_\s*$')
CHECKBOX_RE = re.compile(r'^-\s*\[([ xX])\]\s*(.*)$')


def add_inline_runs(paragraph, text):
    for part in INLINE_TOKEN_RE.split(text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            paragraph.add_run(part)


def _split_into_blocks(text: str) -> list:
    blocks, current = [], []
    for line in text.split('\n'):
        if line.strip() == '':
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(line)
    if current:
        blocks.append(current)
    return blocks


def _classify_block(lines: list):
    first = lines[0].strip()
    if first.startswith('```'):
        inner = lines[1:]
        if inner and inner[-1].strip().startswith('```'):
            inner = inner[:-1]
        meaningful = [l for l in inner if l.strip() and not l.strip().startswith('#')]
        return 'code', (len(meaningful) == 0)
    if first.startswith('|'):
        rows = [l.strip() for l in lines if l.strip().startswith('|')]
        data_rows = rows[2:] if len(rows) >= 2 else []
        has_data = any(any(c.strip() for c in r.strip('|').split('|')[1:]) for r in data_rows)
        return 'table', (not has_data)
    if all(CHECKBOX_RE.match(l.strip()) for l in lines):
        any_checked = any(CHECKBOX_RE.match(l.strip()).group(1).lower() == 'x' for l in lines)
        return 'checkbox', (not any_checked)
    if len(lines) == 1:
        if LABEL_PLACEHOLDER_RE.match(first):
            return 'placeholder', True
        if LABEL_ONLY_RE.match(first):
            return 'label', True
        if first.startswith('> '):
            return 'blockquote', True
        if first.startswith('#### '):
            return 'heading4', False
    return 'text', False


def _render_block(doc, lines: list, kind: str):
    if kind == 'code':
        inner = lines[1:]
        if inner and inner[-1].strip().startswith('```'):
            inner = inner[:-1]
        add_code_block(doc, '\n'.join(inner).strip('\n'))
    elif kind == 'table':
        rows = [l.strip() for l in lines if l.strip().startswith('|')]
        header = [c.strip() for c in rows[0].strip('|').split('|')]
        data = []
        for row_line in rows[2:]:
            cells = [c.strip() for c in row_line.strip('|').split('|')]
            if any(c for c in cells[1:]):
                data.append(cells)
        if data:
            add_table(doc, header, data)
    elif kind == 'checkbox':
        for l in lines:
            m = CHECKBOX_RE.match(l.strip())
            checked, item_text = m.groups()
            if checked.lower() == 'x' and item_text.strip():
                p = doc.add_paragraph(style='List Bullet')
                add_inline_runs(p, '\u2611 ' + item_text.strip())
    elif kind == 'label':
        p = doc.add_paragraph()
        add_inline_runs(p, lines[0].strip())
    elif kind == 'heading4':
        doc.add_heading(lines[0].strip().lstrip('#').strip(), level=4)
    else:
        for line in lines:
            if line.strip().startswith('- ') or line.strip().startswith('* '):
                p = doc.add_paragraph(style='List Bullet')
                add_inline_runs(p, line.strip()[2:].strip())
            else:
                p = doc.add_paragraph()
                add_inline_runs(p, line.strip())


def render_markdown(doc, text: str) -> bool:
    if not text or not text.strip():
        return False
    blocks = _split_into_blocks(text)
    classified = [(b,) + _classify_block(b) for b in blocks]
    rendered_anything = False
    i, n = 0, len(classified)
    while i < n:
        block, kind, empty = classified[i]
        if kind in ('placeholder', 'blockquote'):
            i += 1
            continue
        if kind == 'label':
            if i + 1 < n and not classified[i + 1][2]:
                _render_block(doc, block, kind)
                rendered_anything = True
            i += 1
            continue
        if empty:
            i += 1
            continue
        _render_block(doc, block, kind)
        rendered_anything = True
        i += 1
    return rendered_anything


# ── docx building blocks ────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, htext in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(htext)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr_cells[i], '2B2B2B')
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            add_inline_runs(cells[i].paragraphs[0], str(val))
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    p._p.get_or_add_pPr().append(shd)
    lines = (text or '').split('\n') or ['']
    for i, line in enumerate(lines):
        run = p.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        if i < len(lines) - 1:
            run.add_break()
    return p


def add_placeholder(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x00)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_toc_field(doc):
    """Inserts a real Word TOC field (right-click -> Update Field in Word
    builds it from the document's Heading styles) — matching the official
    template's own mechanism, rather than a hand-built page-number table
    that would immediately go stale."""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    placeholder_text = OxmlElement('w:t')
    placeholder_text.text = "Right-click here and choose \"Update Field\" to generate the table of contents."
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder_text)
    r.append(fld_end)


def enable_auto_update_fields(doc):
    """Sets updateFields in document settings so Word automatically
    refreshes the TOC field (and page-number fields) when the document is
    opened, instead of showing stale placeholder text until someone
    manually right-clicks -> Update Field. This is what actually makes the
    TOC "just work" for anyone opening the file — LibreOffice/Word will
    prompt to update linked fields on open with this set."""
    settings_element = doc.settings.element
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings_element.append(update_fields)


def add_page_number_footer(doc):
    """Matches the official template's footer: a live PAGE field plus a
    'Page' label. Unlike the TOC, page-number fields are computed during
    pagination itself and render correctly even in a static PDF export —
    no manual field update needed."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)
    tail = p.add_run(' | Page')
    tail.font.size = Pt(9)


# ── Title page ───────────────────────────────────────────────────────────────
def build_title_page(doc, candidate, osid, email):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("OffSec Certified Professional")
    run.bold = True
    run.font.size = Pt(26)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Exam Report")
    run.font.size = Pt(16)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run("v2.0")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for value in [email or "[YOUR EMAIL]", f"OSID: {osid}" if osid else "OSID: [YOUR OSID]"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(value)

    if candidate:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run("Candidate: ")
        r1.bold = True
        p.add_run(candidate)

    doc.add_paragraph()
    doc.add_paragraph()
    copyright_p = doc.add_paragraph()
    copyright_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = copyright_p.add_run(
        "Copyright \u00a9 OffSec Ltd. All rights reserved. This report was generated "
        "for personal exam-submission use and follows the structure of OffSec's "
        "official OSCP exam report template."
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()
    doc.add_heading("Table of Contents", level=1)
    add_toc_field(doc)
    doc.add_page_break()


# ── Section 1 ────────────────────────────────────────────────────────────────
def build_section1(doc):
    doc.add_heading("1 OffSec Certified Professional Exam Report", level=1)

    doc.add_heading("1.1 Introduction", level=2)
    doc.add_paragraph(
        "This report documents the results of the OffSec Certified Professional "
        "exam, conducted against the provided exam network. It contains all "
        "efforts undertaken to pass the exam, and is graded on both correctness "
        "and completeness across every required section."
    )

    doc.add_heading("1.2 Objective", level=2)
    doc.add_paragraph(
        "The objective of this assessment is to perform an internal penetration "
        "test against the OffSec exam network, following a methodical approach "
        "from initial access through to final reporting — simulating a real "
        "penetration test engagement start to finish."
    )

    doc.add_heading("1.3 Requirements", level=2)
    doc.add_paragraph("This report includes the following required sections:")
    add_bullets(doc, [
        "Overall High-Level Summary and Recommendations (non-technical)",
        "Methodology walkthrough and detailed outline of steps taken",
        "Each finding with included screenshots, walkthrough, sample code, and proof.txt where applicable",
        "Any additional items not otherwise included",
    ])


# ── Section 2 ────────────────────────────────────────────────────────────────
def build_section2(doc, candidate, results):
    who = candidate or "[Candidate]"
    compromised = [r for r in results if r.get('local') or r.get('proof')]

    doc.add_heading("2 High-Level Summary", level=1)
    add_placeholder(doc,
        f"[Write a non-technical summary here. Who was tested, what the "
        f"objective was, and the overall outcome — e.g. \"{who} was tasked "
        f"with performing an internal penetration test against the OffSec "
        f"exam network. {who} was able to compromise {len(compromised)} of "
        f"{len(results)} target systems, gaining administrative-level access "
        f"where noted below.\" Keep this section free of technical jargon — "
        f"it's meant to be readable by a non-technical stakeholder.]")

    doc.add_heading("2.1 Recommendations", level=2)
    add_placeholder(doc,
        "[General, non-technical remediation guidance — e.g. patching "
        "cadence, credential hygiene, network segmentation. This should be "
        "a short summary; specific fixes belong in each finding's "
        "Vulnerability Fix field below.]")


# ── Section 3 — Methodologies ───────────────────────────────────────────────
def build_section3(doc, candidate, results):
    who = candidate or "[Candidate]"
    compromised = [r for r in results if r.get('local') or r.get('proof')]

    doc.add_heading("3 Methodologies", level=1)
    doc.add_paragraph(
        f"{who} followed a standard penetration testing methodology across "
        f"the exam network, broken out below by phase."
    )

    doc.add_heading("3.1 Information Gathering", level=2)
    doc.add_paragraph(
        "The information gathering phase focused on identifying the scope of "
        "the assessment. The following hosts were in scope:"
    )
    add_bullets(doc, [f"{r['host']}" + (f" ({r['hostname']})" if r.get('hostname') else '') for r in results])

    doc.add_heading("3.2 Service Enumeration", level=2)
    doc.add_paragraph(
        "The service enumeration phase focused on identifying live services "
        "on each in-scope host, to inform potential attack vectors before "
        "attempting exploitation. Full port scan results are included in "
        "each finding below."
    )

    doc.add_heading("3.3 Penetration", level=2)
    doc.add_paragraph(
        f"{who} was able to successfully gain access to {len(compromised)} "
        f"out of {len(results)} systems in scope. Full details of each "
        f"compromised host are included in the following sections."
    )

    doc.add_heading("3.4 Maintaining Access", level=2)
    add_placeholder(doc,
        "[Describe any persistence mechanisms used, if applicable — e.g. "
        "additional accounts created, backdoors installed. Note: many OSCP "
        "exam environments discourage or restrict persistence — only "
        "describe what you actually did.]")

    doc.add_heading("3.5 House Cleaning", level=2)
    add_placeholder(doc,
        "[Describe cleanup performed — removal of any tools, accounts, or "
        "backdoors added during testing. If no persistence was used, note "
        "that explicitly here.]")


# ── Per-finding block (shared by Independent Challenges & AD Set) ──────────
def build_finding_block(doc, number_prefix, r, note_text, flags, include_service_enum=True):
    """Renders one host's findings. number_prefix like '4.1' or '5.2' —
    subsections are numbered off of it (4.1.1, 4.1.2, ...)."""
    sections = parse_note_sections(note_text) if note_text else {}
    attack_notes = sections.get('Attack Notes', '')
    sub = parse_h3_subsections(attack_notes) if attack_notes else {}
    net_access_raw = sections.get('Network Access', '')
    net = parse_network_access(net_access_raw) if net_access_raw else {'access': 'Direct'}

    idx = 1

    # Initial Access — high-level finding writeup (candidate must write the
    # analysis; we can't respell their understanding for them)
    doc.add_heading(f"{number_prefix}.{idx} Initial Access \u2013 [VULNERABILITY NAME]", level=3)
    p = doc.add_paragraph()
    p.add_run("Vulnerability Explanation: ").bold = True
    add_placeholder_inline(doc, "[Describe the root cause of the vulnerability that provided initial access.]")
    p2 = doc.add_paragraph()
    p2.add_run("Vulnerability Fix: ").bold = True
    add_placeholder_inline(doc, "[Describe the remediation.]")
    p3 = doc.add_paragraph()
    p3.add_run("Severity: ").bold = True
    add_placeholder_inline(doc, "[Critical / High / Medium / Low]")
    p4 = doc.add_paragraph()
    p4.add_run("Steps to reproduce the attack: ").bold = True
    add_placeholder_inline(doc, "[Brief narrative summary — the detailed technical walkthrough follows below.]")
    idx += 1

    if is_pivoted(net):
        doc.add_heading(f"{number_prefix}.{idx} Pivoting", level=3)
        pivot_via = net.get('pivot_host') or '[PIVOT HOST]'
        pivot_tool = net.get('pivot_tool') or '[TOOL]'
        doc.add_paragraph(f"This host was reached via a tunnel established through {pivot_via} using {pivot_tool}.")
        if net.get('commands'):
            doc.add_paragraph("Commands used to establish the tunnel:")
            add_code_block(doc, net['commands'])
        else:
            add_placeholder(doc, "[Tunnel command(s) not recorded in notes — fill in manually]")
        idx += 1

    if include_service_enum:
        doc.add_heading(f"{number_prefix}.{idx} Service Enumeration", level=3)
        rows = [[r['host'], ', '.join(r.get('services', [])) or '\u2014']]
        add_table(doc, ["IP Address", "Ports Open"], rows)
        add_placeholder(doc, "[Screenshot: nmap output]")
        idx += 1

    creds = r.get('credentials') or []
    if creds:
        doc.add_heading(f"{number_prefix}.{idx} Credentials", level=3)
        crows = [[c.get('username', ''), c.get('secret', ''), c.get('type', ''),
                   c.get('service', ''), c.get('status', ''), c.get('notes', '')] for c in creds]
        add_table(doc, ["Username", "Secret", "Type", "Service", "Status", "Notes"], crows)
        idx += 1

    # Detailed technical walkthrough — pulled from notes where available.
    # The official template literally reuses "Initial Access" as the
    # heading text for both this and the high-level writeup above (with a
    # different descriptive title each time) — matching that exactly
    # rather than inventing a distinct heading name.
    foothold = sub.get('Foothold', '')
    doc.add_heading(f"{number_prefix}.{idx} Initial Access \u2013 [DETAILED WALKTHROUGH]", level=3)
    rendered = render_markdown(doc, foothold)
    if not rendered:
        add_placeholder(doc, "[Detailed commands/output for initial access \u2014 pull from your note's Foothold section]")
    add_placeholder(doc, "[Screenshot: initial shell / local.txt]")
    p = doc.add_paragraph()
    run = p.add_run(f"local.txt: {flags.get('local', {}).get('hash') or '[HASH]'}")
    run.bold = True
    idx += 1

    privesc = sub.get('Privilege Escalation', '')
    doc.add_heading(f"{number_prefix}.{idx} Privilege Escalation \u2013 [VULNERABILITY NAME]", level=3)
    p = doc.add_paragraph()
    p.add_run("Vulnerability Explanation: ").bold = True
    add_placeholder_inline(doc, "[Describe the root cause of the privilege escalation vector.]")
    p2 = doc.add_paragraph()
    p2.add_run("Vulnerability Fix: ").bold = True
    add_placeholder_inline(doc, "[Describe the remediation.]")
    p3 = doc.add_paragraph()
    p3.add_run("Severity: ").bold = True
    add_placeholder_inline(doc, "[Critical / High / Medium / Low]")
    rendered = render_markdown(doc, privesc)
    if not rendered:
        add_placeholder(doc, "[Detailed commands/output for privilege escalation \u2014 pull from your note's Privilege Escalation section]")
    add_placeholder(doc, "[Screenshot: root/SYSTEM shell + proof.txt]")
    p = doc.add_paragraph()
    run = p.add_run(f"proof.txt: {flags.get('proof', {}).get('hash') or '[HASH]'}")
    run.bold = True
    idx += 1

    # Post-exploitation — lateral movement / loot, if present
    lateral = sub.get('Lateral Movement', '')
    loot = sub.get('Loot', '')
    post_exploit = sub.get('Post Exploitation', '')
    post_body_parts = [b for b in [post_exploit, lateral, loot] if b]
    doc.add_heading(f"{number_prefix}.{idx} Post-Exploitation", level=3)
    rendered_any = False
    for body in post_body_parts:
        if render_markdown(doc, body):
            rendered_any = True
    if not rendered_any:
        add_placeholder(doc, "[Post-exploitation / lateral movement / loot notes, if any]")


def add_placeholder_inline(doc, text):
    """Same as add_placeholder but appends to the paragraph just opened by
    the caller's bold label, instead of starting a new one."""
    p = doc.paragraphs[-1]
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x33, 0x00)


# ── Section 4 — Independent Challenges ──────────────────────────────────────
def build_section4(doc, independent_hosts, notes_by_host, flags_by_host):
    doc.add_heading("4 Independent Challenges", level=1)
    if not independent_hosts:
        doc.add_paragraph("No standalone (non-domain-joined) hosts were in scope for this assessment.")
        return
    for i, r in enumerate(independent_hosts, 1):
        hn = r.get('hostname') or r['host']
        doc.add_heading(f"4.{i} Target #{i} \u2013 {r['host']}" + (f" ({hn})" if r.get('hostname') else ''), level=2)
        build_finding_block(doc, f"4.{i}", r, notes_by_host[r['host']], flags_by_host[r['host']])


# ── Section 5 — Active Directory Set ────────────────────────────────────────
def build_section5(doc, ad_hosts, notes_by_host, flags_by_host):
    doc.add_heading("5 Active Directory Set", level=1)
    if not ad_hosts:
        doc.add_paragraph("No domain-joined hosts were in scope for this assessment.")
        return

    rows = [[r['host'], ', '.join(r.get('services', [])) or '\u2014'] for r in ad_hosts]
    add_table(doc, ["IP Address", "Ports Open"], rows)

    for i, r in enumerate(ad_hosts, 1):
        hn = r.get('hostname') or r['host']
        doc.add_heading(f"5.{i} {hn} \u2013 {r['host']}", level=2)
        build_finding_block(doc, f"5.{i}", r, notes_by_host[r['host']], flags_by_host[r['host']], include_service_enum=False)


# ── Entry point ──────────────────────────────────────────────────────────────
def generate_docx_report(project_dir: Path, include_all: bool = False,
                          candidate: str = '', osid: str = '', email: str = '',
                          exam_date: str = '') -> Document:
    manifest = load_manifest(project_dir)
    all_results = sorted(manifest.values(), key=lambda r: r['host'])
    results = all_results if include_all else [r for r in all_results if r.get('local') or r.get('proof')]
    if not results:
        results = all_results

    results = sorted(results, key=lambda r: access_info(r)[0])

    notes_by_host = {r['host']: load_note(project_dir, r['host']) for r in results}
    sections_by_host = {h: (parse_note_sections(t) if t else {}) for h, t in notes_by_host.items()}
    flags_by_host = {h: parse_flags(s) for h, s in sections_by_host.items()}
    domain_by_host = {h: parse_domain(s) for h, s in sections_by_host.items()}

    independent_hosts = [r for r in results if not is_ad_domain_joined(domain_by_host[r['host']])]
    ad_hosts = [r for r in results if is_ad_domain_joined(domain_by_host[r['host']])]

    doc = Document()
    add_page_number_footer(doc)
    enable_auto_update_fields(doc)
    build_title_page(doc, candidate, osid, email)
    build_section1(doc)
    build_section2(doc, candidate, results)
    build_section3(doc, candidate, results)
    build_section4(doc, independent_hosts, notes_by_host, flags_by_host)
    build_section5(doc, ad_hosts, notes_by_host, flags_by_host)
    return doc


def parse_args():
    p = argparse.ArgumentParser(description="Generate an OSCP-exam-format .docx report from a recon_scan.py project")
    p.add_argument('project', help='Project directory')
    p.add_argument('-o', '--output', help='Output path (default: <project>/FINAL_REPORT.docx)')
    p.add_argument('--all', action='store_true', help='Include every scanned host, not just Local/Proof-flagged ones')
    p.add_argument('--candidate', default='', help='Candidate name')
    p.add_argument('--osid', default='', help='OSID')
    p.add_argument('--email', default='', help='Email')
    p.add_argument('--exam-date', default='', help='(unused — the official template has no exam-date field)')
    return p.parse_args()


def main():
    args = parse_args()
    project_dir = Path(args.project)
    if not project_dir.exists():
        print(f"[-] Project directory not found: {project_dir}", file=sys.stderr)
        sys.exit(1)

    doc = generate_docx_report(
        project_dir, include_all=args.all,
        candidate=args.candidate, osid=args.osid, email=args.email,
    )
    out_path = Path(args.output) if args.output else project_dir / 'FINAL_REPORT.docx'
    doc.save(str(out_path))
    print(f"[+] Report written: {out_path}")


if __name__ == '__main__':
    main()
